# -*- coding: utf-8 -*-
"""
生成《证书生成及系统服务协议》Word 文档。
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DOCX_PATH = os.path.join(os.path.dirname(__file__), "证书生成及系统服务协议.docx")

CN_FONT = "宋体"
TITLE_FONT = "黑体"


def set_cn_font(run, font=CN_FONT, size=10.5, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_cn_font(run, font=TITLE_FONT, size=18, bold=True)
    p.paragraph_format.space_after = Pt(6)


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, font=TITLE_FONT, size=13, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_para(doc, segments, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, indent=False):
    """segments: 字符串，或 (text, bold) 元组列表"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if isinstance(segments, str):
        segments = [(segments, False)]
    for text, bold in segments:
        run = p.add_run(text)
        set_cn_font(run, size=size, bold=bold)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("—" * 40)
    set_cn_font(run, size=9, color=RGBColor(0xBB, 0xBB, 0xBB))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)


def add_grid_table(doc, headers, rows_data, widths=None):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_cn_font(run, size=10, bold=True)
    for i, row in enumerate(rows_data, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            run = cell.paragraphs[0].add_run(val)
            set_cn_font(run, size=10)
    return table


def build_attachments(doc):
    # ---------- 附件一 ----------
    doc.add_page_break()
    add_heading(doc, "附件一  附赠工具功能范围与使用授权清单")
    add_para(doc, [("附赠工具名称：", True), ("AIGC 大赛评分工具", False)])
    add_para(doc, [("授权性质：", True), ("非独占、不可转让、仅限甲方内部使用", False)])
    add_para(doc, [("授权期限：", True), ("自本协议生效之日起至 ______ 年 ____ 月 ____ 日止", False)])
    add_para(doc, [("使用账号数量：", True), ("________ 个", False)])
    add_para(doc, [("1. 功能范围清单", True)])
    add_grid_table(
        doc,
        ["序号", "功能模块", "功能说明", "是否开放"],
        [
            ["1", "作品录入", "支持文本/图片等 AIGC 作品的批量录入", "是 □ / 否 □"],
            ["2", "智能评分", "基于预设规则/模型对作品自动打分", "是 □ / 否 □"],
            ["3", "评分维度配置", "自定义评分维度及权重", "是 □ / 否 □"],
            ["4", "结果导出", "评分结果导出为 Excel/CSV", "是 □ / 否 □"],
            ["5", "排名统计", "按得分生成排名及统计报表", "是 □ / 否 □"],
        ],
    )
    add_para(doc, [("2. 使用限制", True)])
    add_para(doc, "本附赠工具的使用限制遵循本协议第五条第 5.3 条之约定，甲方不得超出本清单范围使用。")
    add_para(doc, [("3. 技术支持", True)])
    add_para(doc, "乙方提供附赠工具的基础使用指导；超出基础范围的定制开发、二次集成及运维服务需另行协商。")

    # ---------- 附件二 ----------
    doc.add_page_break()
    add_heading(doc, "附件二  证书模板及数据格式确认单")
    add_para(doc, [("1. 证书模板信息", True)])
    add_grid_table(
        doc,
        ["项目", "确认内容"],
        [
            ["模板名称", "__________________________"],
            ["模板尺寸（像素）", "宽 ________ × 高 ________"],
            ["文件格式", "PNG □ / JPG □ / PDF □"],
            ["模板提供方", "甲方 □ / 乙方代设计 □"],
            ["确认日期", "______ 年 ____ 月 ____ 日"],
        ],
    )
    add_para(doc, [("2. 数据字段映射", True)])
    add_para(doc, "甲方提供的 Excel 数据列与证书模板占位字段的对应关系如下：")
    add_grid_table(
        doc,
        ["序号", "证书占位字段", "Excel 列名", "示例值", "是否必填"],
        [
            ["1", "姓名", "__________", "张三", "是"],
            ["2", "证书编号", "__________", "NO.2026001", "是"],
            ["3", "颁发日期", "__________", "2026-06-05", "是"],
            ["4", "________", "__________", "__________", "是 □ / 否 □"],
            ["5", "________", "__________", "__________", "是 □ / 否 □"],
        ],
    )
    add_para(doc, "说明：Excel 首行为表头，自第二行起为数据；字段名称应与证书模板占位字段严格一致。")

    # ---------- 附件三 ----------
    doc.add_page_break()
    add_heading(doc, "附件三  交付与验收标准")
    add_para(doc, [("1. 交付内容与方式", True)])
    add_grid_table(
        doc,
        ["项目", "约定内容"],
        [
            ["证书总数量", "15,000 份（首批）"],
            ["交付格式", "PNG □ / PDF □"],
            ["命名规则", "__________________________"],
            ["交付方式", "上传小程序后台 □ / 网盘 □ / API □"],
            ["交付期限", "自数据与模板确认之日起 ____ 个工作日内"],
        ],
    )
    add_para(doc, [("2. 验收标准", True)])
    for t in [
        "2.1 数量：生成并成功上传的证书数量与甲方提供的有效数据记录数一致；",
        "2.2 准确性：证书上的姓名、编号、日期等字段与基础数据一致，无错字、错位；",
        "2.3 清晰度：证书图片清晰，文字无遮挡、无变形，排版符合模板设计；",
        "2.4 可访问性：上传至小程序后台的证书可正常读取、下载及展示。",
    ]:
        add_para(doc, t)
    add_para(doc, [("3. 验收流程", True)])
    for t in [
        "3.1 乙方完成交付后以书面（含数据电文）形式通知甲方验收；",
        "3.2 甲方应在 ____ 个工作日内完成验收并出具书面意见；逐期未提出书面异议的，视为验收合格；",
        "3.3 验收不合格的，甲方应明确列明问题清单，乙方在 ____ 个工作日内进行修正后重新提交验收；因甲方数据/模板问题导致的返工除外。",
    ]:
        add_para(doc, t)


def build():
    doc = Document()
    # 默认正文样式
    style = doc.styles["Normal"]
    style.font.name = CN_FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    # 标题
    add_title(doc, "证书生成及系统服务协议")
    add_para(doc, [("合同编号：__________________", True)], align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 双方信息
    add_para(doc, [("甲方（委托方）：__________________________", True)])
    for t in [
        "统一社会信用代码：__________________________",
        "法定代表人：____________________",
        "注册地址：______________________________",
        "联系电话：____________________",
    ]:
        add_para(doc, "    " + t)

    add_para(doc, [("乙方（服务提供方 / 个人）：____________________", True)])
    for t in [
        "身份证号码：__________________________",
        "联系地址：______________________________",
        "联系电话：____________________",
    ]:
        add_para(doc, "    " + t)

    add_para(doc, [("说明：甲方为委托生成证书的采购方；乙方为以个人身份提供证书生成及系统服务的自然人。", False)], size=9)

    add_divider(doc)

    # 鉴于条款
    add_heading(doc, "鉴于条款（前言）")
    for t in [
        "1. 乙方系具备证书批量生成软件开发及相关数字化服务能力的自然人，对其证书生成系统及附赠工具拥有相应的技术能力与知识产权；",
        "2. 甲方因业务需要，拟委托乙方批量生成电子证书并对接甲方指定的微信小程序后台系统；",
        "3. 双方本着平等自愿、诚实信用、互利共赢的原则，根据《中华人民共和国民法典》及相关法律法规，经友好协商，就本次服务事宜达成如下协议，以资共同遵守。",
    ]:
        add_para(doc, t)

    add_divider(doc)

    # 第一条
    add_heading(doc, "第一条  定义与解释")
    add_para(doc, "除本协议另有约定外，下列术语具有如下含义：")
    defs = [
        ('1.1 "证书"：', "指乙方依据甲方提供的数据信息，通过乙方证书批量生成系统在指定模板基础上渲染生成的电子证书图片及/或 PDF 文件。"),
        ('1.2 "证书模板"：', "指甲方提供或经甲方书面确认的、用于生成证书的空白底图及版式设计。"),
        ('1.3 "基础数据"：', "指甲方为生成证书而向乙方提供的、包含姓名、编号、日期等字段的 Excel 等格式的结构化数据。"),
        ('1.4 "小程序后台"：', "指甲方指定的、用于接收和存储证书文件的微信小程序服务端管理系统。"),
        ('1.5 "AIGC 大赛评分工具"（以下简称"附赠工具"）：', "指乙方自主研发、作为本次合作附赠项提供甲方使用的人工智能生成内容评分软件工具。"),
        ('1.6 "交付成果"：', "指本协议项下乙方应向甲方交付的全部证书文件及相关服务成果。"),
    ]
    for head, body in defs:
        add_para(doc, [(head, True), (body, False)])

    add_divider(doc)

    # 第二条
    add_heading(doc, "第二条  服务内容")
    add_para(doc, [("2.1 证书生成服务", True)])
    add_para(doc, "2.1.1 乙方负责依据甲方确认的证书模板及提供的基础数据，批量生成 15,000（壹万伍仟）份 电子证书。")
    add_para(doc, "2.1.2 证书输出格式为 ______（PNG 图片 / PDF 文件 / 两者皆有，由双方在附件中确认）。")
    add_para(doc, "2.1.3 乙方负责将生成的全部证书文件按甲方要求的命名规则及目录结构，上传至甲方指定的微信小程序后台，或按双方约定的方式（如网盘、API 接口等）交付。")
    add_para(doc, [("2.2 附赠工具使用权", True)])
    add_para(doc, "2.2.1 作为本次合作的附赠服务，乙方授予甲方在本协议有效期内非独占、不可转让的 AIGC 大赛评分工具使用权，供甲方内部业务使用。")
    add_para(doc, "2.2.2 附赠工具的具体功能范围、使用账号数量及使用期限见本协议附件一，附赠工具的权利归属及限制详见本协议第五条。")
    add_para(doc, [("2.3 服务范围之外", True)])
    add_para(doc, "本协议服务内容以本条及附件明确列举的事项为限。甲方提出的、超出本协议约定范围的任何需求（包括但不限于模板重新设计、数据清洗、额外系统开发、定制化功能等），双方应另行协商并签订补充协议，相关费用另计。")

    add_divider(doc)

    # 第三条
    add_heading(doc, "第三条  费用与结算")
    add_para(doc, [("3.1 服务总费用", True)])
    add_para(doc, [("本协议项下首批 15,000 份证书生成及上传服务、附赠工具使用权的服务总费用为人民币 ", False),
                   ("3,000 元（大写：叁仟元整）", True), ("。", False)])
    add_para(doc, [("3.2 超额证书阶梯计费", True)])
    add_para(doc, "首批 15,000 份证书完成后，若甲方需要追加生成证书，超出部分按以下阶梯单价计费：")

    # 阶梯表格
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["阶梯档位", "累计追加数量区间（份）", "单价（人民币元/份）"]
    rows_data = [
        ["第一档", "1 — 5,000", "__________"],
        ["第二档", "5,001 — 20,000", "__________"],
        ["第三档", "20,001 份及以上", "__________"],
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_cn_font(run, size=10, bold=True)
    for i, row in enumerate(rows_data, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(val)
            set_cn_font(run, size=10)

    add_para(doc, '3.2.1 阶梯计费按"分段累进"方式计算，即各区间数量分别适用对应档位单价后加总。')
    add_para(doc, "3.2.2 追加证书需求由甲方以书面（含电子邮件、数据电文）形式提出并经乙方书面确认后执行，费用按确认数量结算。")
    add_para(doc, [("3.3 支付方式与节点", True)])
    add_para(doc, "3.3.1 预付款：本协议签订之日起 ____ 个工作日内，甲方支付总费用的 ____%，即人民币 __________ 元；乙方收款后启动服务。")
    add_para(doc, "3.3.2 尾款：乙方完成全部证书上传交付并经甲方验收合格后 ____ 个工作日内，甲方支付剩余 ____%，即人民币 __________ 元。")
    add_para(doc, "3.3.3 追加费用：超额证书费用于每批次交付验收合格后 ____ 个工作日内一次性结清。")
    add_para(doc, [("3.4 付款方式：", True), ("甲方通过银行转账方式支付至乙方以下个人账户：", False)])
    for t in ["账户姓名：__________________", "开户银行：__________________", "银行账号：__________________"]:
        add_para(doc, "    " + t)
    add_para(doc, [("3.5 税费与票据：", True), ("本协议费用为乙方实得净额。如甲方需要发票，由乙方协助至主管税务机关代开，因代开发票所产生的税费及甲方依法应履行的代扣代缴义务，由双方依照法律规定及__________（甲方/乙方）承担。", False)])
    add_para(doc, [("3.6 逾期付款：", True), ("甲方未按约定支付任何款项的，每逾期一日，应按逾期金额的 万分之五（0.05%） 向乙方支付违约金；逾期超过 ____ 日的，乙方有权暂停服务且不视为乙方违约。", False)])

    add_divider(doc)

    # 第四条
    add_heading(doc, "第四条  双方权利与义务")
    add_para(doc, [("4.1 乙方（服务提供方）权利与义务", True)])
    for t in [
        "4.1.1 乙方应按本协议约定的标准、格式与时间完成证书生成及上传服务；",
        "4.1.2 乙方有权要求甲方及时、完整、准确地提供基础数据及证书模板，并提供必要的系统对接配合；",
        "4.1.3 因甲方提供的数据或模板存在错误、缺失、格式不符导致的返工，乙方有权要求甲方承担相应额外费用及顺延工期；",
        "4.1.4 乙方应对在服务过程中获知的甲方数据及商业信息承担保密义务（详见第六条）；",
        "4.1.5 乙方有权在甲方逾期付款或严重违约时暂停或终止服务。",
    ]:
        add_para(doc, t)
    add_para(doc, [("4.2 甲方（委托方）权利与义务", True)])
    for t in [
        "4.2.1 甲方应在本协议签订后 ____ 个工作日内，向乙方提供符合要求的证书模板及基础数据，并对数据的真实性、合法性、准确性及完整性承担全部责任；",
        "4.2.2 甲方应保证其提供的基础数据已取得相关数据主体的合法授权，不侵犯任何第三方的合法权益；因数据来源或内容引发的纠纷、行政处罚或第三方索赔，由甲方独立承担，乙方不承担任何责任；",
        "4.2.3 甲方应提供小程序后台的对接接口、账号权限及技术文档，并指定专人配合系统对接调试；因甲方系统接口故障、权限未开通或配合不及时导致的工期延误，不视为乙方违约；",
        "4.2.4 甲方应按本协议约定及时支付各项费用；",
        "4.2.5 甲方应在收到交付成果后 ____ 个工作日内完成验收。逾期未提出书面异议的，视为验收合格。",
    ]:
        add_para(doc, t)
    add_para(doc, [("4.3 数据交付与对接", True)])
    add_para(doc, "4.3.1 双方应指定专门的项目联系人，负责数据传递、系统对接及问题沟通；")
    add_para(doc, "4.3.2 基础数据的传递应采用安全方式（如加密传输），双方均应妥善保管，防止泄露。")

    add_divider(doc)

    # 第五条
    add_heading(doc, "第五条  知识产权")
    add_para(doc, [("5.1 乙方系统及工具的知识产权归属", True)])
    add_para(doc, "5.1.1 乙方用于提供本服务的证书生成系统、软件、源代码、技术文档及附赠的 AIGC 大赛评分工具，其全部知识产权（包括但不限于著作权、专利权、商业秘密等）均归乙方或其合法权利人单独所有，本协议的签订与履行不构成上述任何知识产权的转让。")
    add_para(doc, [("5.2 证书成果的权利归属", True)])
    add_para(doc, "5.2.1 乙方依据甲方数据生成的最终证书文件，其使用权归甲方所有；甲方提供的模板、数据、品牌标识等内容的知识产权仍归甲方或其合法权利人所有。")
    add_para(doc, [("5.3 附赠工具的使用限制", True)])
    add_para(doc, "甲方对附赠工具的使用，应遵守以下限制条款：")
    for t in [
        "5.3.1 甲方仅享有约定范围内的使用权，不享有所有权，不得据此主张任何知识产权；",
        "5.3.2 甲方不得对附赠工具进行复制、出售、出租、转授权、分发或以任何形式向第三方提供；",
        "5.3.3 甲方不得对附赠工具进行反向工程、反编译、反汇编，或以其他方式试图获取其源代码、算法或底层逻辑；",
        "5.3.4 甲方不得删除、修改附赠工具中的任何权利声明、商标标识或技术保护措施；",
        '5.3.5 附赠工具按"现状（AS-IS）"提供，乙方不对其适用于甲方特定目的作出任何明示或默示担保；除乙方故意或重大过失外，乙方不对甲方使用附赠工具产生的任何后果承担责任；',
        "5.3.6 甲方违反本条任何约定的，乙方有权立即终止授权并要求甲方承担违约责任及赔偿由此造成的全部损失。",
    ]:
        add_para(doc, t)

    add_divider(doc)

    # 第六条
    add_heading(doc, "第六条  保密条款")
    add_para(doc, "6.1 双方对因履行本协议而知悉的对方商业秘密、技术信息、数据资料及其他保密信息，均负有保密义务，未经对方书面同意，不得向任何第三方披露或用于本协议之外的目的。")
    add_para(doc, "6.2 本条保密义务在本协议终止或解除后继续有效，期限为 ____ 年。")
    add_para(doc, "6.3 下列信息不属于保密信息：（1）已进入公共领域的信息；（2）接收方在收到前已合法持有的信息；（3）依法律法规或有权机关要求必须披露的信息（但应在法律允许范围内事先通知对方）。")

    add_divider(doc)

    # 第七条
    add_heading(doc, "第七条  违约责任")
    add_para(doc, "7.1 任何一方未履行或未完全履行本协议约定义务的，应承担继续履行、采取补救措施或赔偿损失等违约责任。")
    add_para(doc, [("7.2 乙方违约：", True), ("乙方无正当理由未按期完成证书交付的，每逾期一日，按服务总费用的 万分之三（0.03%） 向甲方支付违约金；但因甲方原因（数据/模板缺陷、对接不配合、付款逾期等）导致的延误除外。", False)])
    add_para(doc, [("7.3 甲方违约：", True)])
    add_para(doc, "7.3.1 甲方逾期付款的，按第 3.6 条承担违约责任；")
    add_para(doc, "7.3.2 甲方违反第五条知识产权及使用限制条款的，应向乙方支付违约金人民币 __________ 元；违约金不足以弥补乙方损失的，甲方应就不足部分继续赔偿。")
    add_para(doc, [("7.4 赔偿范围限制：", True), ("除因一方故意或重大过失、以及侵犯知识产权、违反保密义务的情形外，任何一方依本协议向对方承担的赔偿责任总额，不超过本协议服务总费用；且任何一方均不对间接损失、利润损失等承担责任。", False)])
    add_para(doc, [("7.5 不可抗力：", True), ("因不可抗力（包括但不限于自然灾害、政府行为、网络基础设施重大故障、第三方平台（如微信）政策变更等）导致无法履行的，遭遇方应及时通知并提供证明，可相应免责或顺延履行期限。", False)])

    add_divider(doc)

    # 第八条
    add_heading(doc, "第八条  协议的变更、解除与终止")
    add_para(doc, "8.1 本协议的任何变更、补充均须经双方书面协商一致，签订书面补充协议，补充协议与本协议具有同等法律效力。")
    add_para(doc, "8.2 有下列情形之一的，守约方有权书面通知解除本协议：")
    for t in [
        "（1）一方根本违约，致使协议目的无法实现；",
        "（2）一方逾期付款 / 逾期交付超过 ____ 日，经催告后仍未履行；",
        "（3）法律法规规定的其他情形。",
    ]:
        add_para(doc, "    " + t)
    add_para(doc, "8.3 本协议解除或终止的，不影响第五条（知识产权）、第六条（保密）、第七条（违约责任）、第九条（争议解决）等条款的继续有效，亦不影响守约方追究违约方责任的权利。")

    add_divider(doc)

    # 第九条
    add_heading(doc, "第九条  法律适用与争议解决")
    add_para(doc, "9.1 本协议的订立、效力、解释、履行及争议解决，均适用中华人民共和国法律（为本协议之目的，不包括香港特别行政区、澳门特别行政区及台湾地区法律）。")
    add_para(doc, [("9.2 因本协议引起的或与本协议有关的任何争议，双方应首先通过友好协商解决；协商不成的，任何一方均有权向 ", False),
                   ("乙方住所地有管辖权的人民法院", True), (" 提起诉讼。", False)])
    add_para(doc, "9.3 诉讼期间，除争议部分外，双方应继续履行本协议其他条款。")

    add_divider(doc)

    # 第十条
    add_heading(doc, "第十条  其他")
    add_para(doc, "10.1 本协议自甲方加盖公章（或授权代表签字）、乙方本人签字之日起生效。")
    add_para(doc, "10.2 本协议未尽事宜，由双方另行协商并签订补充协议。")
    add_para(doc, "10.3 本协议下列附件为本协议不可分割的组成部分，与本协议具有同等法律效力：")
    for t in [
        "附件一：附赠工具功能范围与使用授权清单",
        "附件二：证书模板及数据格式确认单",
        "附件三：交付与验收标准",
    ]:
        add_para(doc, "    " + t)
    add_para(doc, "10.4 本协议一式 ____ 份，甲乙双方各执 ____ 份，具有同等法律效力。")
    add_para(doc, "10.5 双方通过电子邮件、即时通讯等数据电文形式确认的需求、数据及验收意见，作为本协议履行的有效凭证。")

    add_divider(doc)

    # 签署页
    add_para(doc, [("（以下无正文，为签署页）", True)], align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    sign = doc.add_table(rows=4, cols=2)
    sign.style = "Table Grid"
    sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_data = [
        ["甲方（委托方）：（盖章）", "乙方（服务提供方）：（签字）"],
        ["法定代表人 / 授权代表（签字）：\n\n______________________________", "本人签字：\n\n______________________________"],
        ["身份证号 / 统一社会信用代码：\n\n______________________________", "身份证号：\n\n______________________________"],
        ["日期：______年____月____日", "日期：______年____月____日"],
    ]
    for i, row in enumerate(sign_data):
        for j, val in enumerate(row):
            cell = sign.rows[i].cells[j]
            cell.paragraphs[0].text = ""
            for k, line in enumerate(val.split("\n")):
                para = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                run = para.add_run(line)
                set_cn_font(run, size=10)

    doc.add_paragraph()
    tip = add_para(doc, [("法律提示：", True),
                         ("本文本为商务谈判用合同范本，其中带下划线的空格、违约金比例、计费单价、税费承担及管辖约定等关键条款需双方据实填写。乙方以个人身份签约，建议留存身份证复印件、收款记录及沟通凭证；正式签署前，建议交由执业律师结合具体交易背景及最新法律法规进行终审，以充分保障己方权益。", False)],
                        size=9)

    # 附件
    build_attachments(doc)

    doc.save(DOCX_PATH)
    print("已生成:", DOCX_PATH)


if __name__ == "__main__":
    build()
